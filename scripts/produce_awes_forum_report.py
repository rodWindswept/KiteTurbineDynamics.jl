"""
TRPT AWE Forum Report v3 — Generator
=====================================
Produces TRPT_AWE_Forum_Report_v3.docx from committed figures and v4/v5
campaign data.  Uses python-docx to match the style of other KTD reports.

Usage:
    python3 scripts/produce_awes_forum_report.py

Output:
    TRPT_AWE_Forum_Report_v3.docx   (repo root)
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import csv, datetime

ROOT    = Path(__file__).parent.parent
FIGS    = ROOT / "figures" / "report"
OUT     = ROOT / "TRPT_AWE_Forum_Report_v3.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x2A)
TEAL   = RGBColor(0x00, 0x7A, 0x87)
SLATE  = RGBColor(0x44, 0x4F, 0x5A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF4, 0xF8)
MID    = RGBColor(0xD6, 0xE4, 0xED)
GREEN  = RGBColor(0x2B, 0x8A, 0x3E)
AMBER  = RGBColor(0xC7, 0x70, 0x00)
RED    = RGBColor(0xB4, 0x2D, 0x2D)

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def cell_font(cell, bold=False, color=SLATE, size=10, align=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name  = "Calibri"
            run.font.size  = Pt(size)
            run.font.color.rgb = color
            run.bold       = bold
        if align is not None:
            para.alignment = align

def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(level=level)
    p.clear()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.color.rgb = NAVY
    r.bold = True
    r.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    return p

def para(doc: Document, text: str, bold=False, italic=False,
         color=SLATE, size=11, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(size)
    r.font.color.rgb = color
    r.bold   = bold
    r.italic = italic
    if align is not None:
        p.alignment = align
    pPr = p.paragraph_format
    pPr.space_before = Pt(space_before)
    pPr.space_after  = Pt(space_after)
    return p

def caption(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(9)
    r.font.color.rgb = SLATE
    r.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    return p

def add_figure(doc: Document, filename: str, cap: str,
               width_in: float = 5.5):
    path = FIGS / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(path), width=Inches(width_in))
    else:
        para(doc, f"[Figure not found: {filename}]", italic=True, color=RED)
    caption(doc, cap)

def hline(doc: Document, color_hex="D6E4ED", thickness=6):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thickness))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(6)

def summary_table(doc: Document, headers: list, rows: list,
                  col_widths_in: list | None = None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Set column widths if provided
    if col_widths_in:
        for i, cell in enumerate(tbl.rows[0].cells):
            cell.width = Inches(col_widths_in[i])

    # Header row
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        set_cell_bg(cell, NAVY)
        cell.text = h
        cell_font(cell, bold=True, color=WHITE, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for i, row in enumerate(rows):
        bg = LIGHT if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            cell = tbl.rows[i+1].cells[j]
            set_cell_bg(cell, bg)
            cell.text = str(val)
            cell_font(cell, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    return tbl


# ── Campaign data ──────────────────────────────────────────────────────────────
CAMPAIGN_PROGRESSION = [
    ("v2", "Apr 2026", "2.808 kg", "19.22 kg",
     "Euler buckling only — torsionally infeasible (54/60 islands fail)"),
    ("v3", "Apr 2026", "15.435 kg", "145.88 kg",
     "Torsional FOS ≥ 1.5 added; cylindrical geometry forced by uniform spacing"),
    ("v4", "Apr 2026", "10.587 kg", "79.51 kg",
     "Geometric L/r ring spacing; taper restored as design variable (−31 % vs v3)"),
    ("v5", "Apr 2026", "11.470 kg", "39.30 kg",
     "BEM-coupled rotor radius (self-consistent R from n_lines/Cp); 50 kW −51 % vs v4"),
]

V5_10KW_WINNER = {
    "n_lines":        "8",
    "mass_kg":        "11.470 kg",
    "r_hub_m":        "1.600 m",
    "r_bottom_m":     "0.336 m",
    "taper_ratio":    "0.210",
    "target_Lr":      "2.00",
    "tether_length":  "30.0 m",
    "Do_top_m":       "40.9 mm",
    "t_over_D":       "0.020",
    "min_fos":        "1.800 (Euler + torsional)",
}

V5_50KW_WINNER = {
    "n_lines":        "8",
    "mass_kg":        "39.295 kg",
    "r_hub_m":        "3.578 m",
    "r_bottom_m":     "0.300 m",
    "tether_length":  "67.08 m",
    "Do_top_m":       "58.6 mm",
    "min_fos":        "1.800 (Euler + torsional)",
}

V6_OPEN_QUESTIONS = [
    ("CFD / panel-method validation of n_lines = 8",
     "Strip theory is not validated above n = 6 lines. Blade-to-blade wake "
     "interference and solidity blockage are unmodelled. This is the highest-priority "
     "item before any hardware commitment."),
    ("Joint β + structural optimisation",
     "Elevation angle β was fixed at 30° throughout v2–v5. Cold-start and "
     "lift-kite analysis suggest the optimum β is near 26°. v6 should free β "
     "alongside the structural parameters to find the true aerodynamic-structural "
     "optimum simultaneously."),
    ("Dynamic torsional loading and fatigue",
     "All campaigns size against a static peak load envelope with Design Load "
     "Factor DLF = 1.2. Cyclic 1P/2P tether tension loading and S-N fatigue "
     "are not modelled. Fatigue life must be assessed before selecting tube "
     "wall thickness for manufacture."),
]


# ── Document build ─────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Cover ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("TRPT Kite Turbine")
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("Structural Optimisation — Campaigns v2 to v5")
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.color.rgb = TEAL
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("AWE Forum Technical Report  ·  v3  ·  April 2026")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = SLATE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_figure(doc, "fig_trpt_system.png",
               "Figure 1 — TRPT system overview: helical tether lines transmit torque "
               "from the airborne rotor to the ground generator.", width_in=5.0)

    doc.add_page_break()

    # ── 1. Executive summary ───────────────────────────────────────────────────
    heading(doc, "1  Executive Summary")
    hline(doc)
    para(doc,
         "This report presents structural optimisation results for the Tensile Rotary "
         "Power Transmission (TRPT) kite turbine shaft across four Differential Evolution "
         "campaigns (v2–v5, April 2026).  Each campaign added one or more physical "
         "constraints or aerodynamic models to the fitness function, progressively closing "
         "the gap between the optimised design and physical reality.")
    para(doc,
         "The key headline results are:")

    bullets = [
        "n_lines = 8 tether lines selected unanimously across all 120 islands (v4 + v5), "
        "for both 10 kW and 50 kW configurations.",
        "10 kW shaft minimum mass: 11.470 kg (v5, BEM-coupled), down from 15.435 kg "
        "in v3 after adding the torsional constraint — a 25.7 % reduction.",
        "50 kW shaft minimum mass: 39.295 kg (v5), a 51 % reduction vs v4 (79.5 kg), "
        "driven by more accurate aerodynamic loading from the BEM model.",
        "Geometric L/r ring spacing (v4) is the key enabling innovation: it allows a "
        "tapered shaft geometry without the Euler buckling penalty that forced cylindrical "
        "designs in v3.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(b)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.color.rgb = SLATE

    doc.add_paragraph()
    para(doc,
         "Three open questions remain before hardware commitment: CFD validation of "
         "n_lines = 8 aerodynamics (strip theory not validated above n = 6), joint "
         "β + structural optimisation (β fixed at 30° throughout), and dynamic fatigue "
         "modelling (currently static peak envelope only).")

    doc.add_page_break()

    # ── 2. System overview ─────────────────────────────────────────────────────
    heading(doc, "2  System Overview")
    hline(doc)
    para(doc,
         "The TRPT transmits power from an airborne rotor to a ground-level generator "
         "through a tensile column of rings connected by helical tether lines.  The shaft "
         "is supported by tension in the tether lines — there are no rigid compression "
         "members.  Torque propagates through the helical geometry: twist of the shaft "
         "increases line tension asymmetrically, generating a net restoring torque.")

    para(doc,
         "The canonical system parameters used throughout v2–v5 are:")

    summary_table(doc,
        ["Parameter", "Value"],
        [
            ["Rated power",           "10 kW / 50 kW (separate optimisations)"],
            ["Rated wind speed",      "11 m/s"],
            ["Tether length L",       "30 m (10 kW) / 67 m (50 kW)"],
            ["Elevation angle β",     "30° (fixed throughout v2–v5)"],
            ["Rotor radius R",        "5.0 m (v2–v4) / 5.12 m (v5 BEM-derived)"],
            ["n_lines (canonical)",   "5 (baseline) → 8 (optimised)"],
            ["Design Load Factor",    "1.2 (calibrated from ODE simulation)"],
            ["Euler FOS requirement", "≥ 1.8"],
            ["Torsional FOS requirement", "≥ 1.5 (v3 onwards, Tulloch/Wacker)"],
        ],
        col_widths_in=[2.8, 4.0])

    doc.add_page_break()

    # ── 3. Campaign progression ────────────────────────────────────────────────
    heading(doc, "3  Campaign Progression")
    hline(doc)
    para(doc,
         "Four optimisation campaigns were run sequentially, each building on the "
         "structural model and results of the previous.  60 Differential Evolution "
         "islands (2 power configs × 3 beam profiles × 5 parameter zones × 2 RNG seeds) "
         "were run per campaign.")

    add_figure(doc, "fig_campaign_progression.png",
               "Figure 2 — Minimum shaft mass across campaigns v2–v5 for 10 kW (blue) "
               "and 50 kW (orange).  The v2 designs are physically infeasible — their "
               "apparent low mass is due to the absent torsional constraint.",
               width_in=5.5)

    heading(doc, "3.1  Campaign summary table", level=2)
    summary_table(doc,
        ["Campaign", "Date", "10 kW winner", "50 kW winner", "Key change"],
        CAMPAIGN_PROGRESSION,
        col_widths_in=[1.0, 0.9, 1.2, 1.2, 3.4])

    # v2
    heading(doc, "3.2  v2 — Euler buckling only (baseline)", level=2)
    para(doc,
         "v2 used 12 degrees of freedom and enforced only Euler buckling (FOS ≥ 1.8).  "
         "The apparent winner at 2.808 kg is physically invalid: post-hoc torsional "
         "analysis showed 54 of 60 island winners fail the Tulloch/Wacker torsional "
         "stability criterion, with the global winner having torsional FOS = 0.069 — "
         "failing by 22×.  The dominant failure mode is small ring radius at the ground "
         "end: torsional capacity scales as r², so a tapered shaft with r_bottom ≈ 0.1 m "
         "has negligible torque capacity.")

    # v3
    heading(doc, "3.3  v3 — Torsional constraint added", level=2)
    para(doc,
         "v3 added the Tulloch/Wacker torsional stability criterion as a hard feasibility "
         "gate.  The maximum torque capacity of a segment of axial length L between rings "
         "of radius r under total tether tension T_total is:")
    p = doc.add_paragraph()
    r = p.add_run("τ_cap = T_total × r² / √(L² + 2r²)")
    r.font.name = "Courier New"
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc,
         "All 60 v3 islands converged to taper_ratio = 1.0 (cylindrical shaft).  With "
         "uniform ring spacing, a tapered shaft has high L/r at the thin end, creating "
         "an Euler buckling limit that forces cylindrical geometry.  The v3 10 kW winner "
         "is 15.435 kg — physically valid, but the cylindrical constraint is artificial.")

    add_figure(doc, "fig_tulloch_wacker_chart.png",
               "Figure 3 — Tulloch/Wacker torsional stability criterion.  "
               "τ_cap ∝ r² / √(L² + 2r²).  Small ring radius is the dominant failure mode.",
               width_in=5.0)

    # v4
    heading(doc, "3.4  v4 — Geometric L/r ring spacing", level=2)
    para(doc,
         "The key insight from v3 is that uniform ring spacing penalises taper.  v4 "
         "replaced uniform spacing with geometric-series spacing that maintains a "
         "constant L/r ratio per segment (target_Lr parameter, 9th degree of freedom).  "
         "This allows the optimiser to use a tapered shaft — larger hub radius, smaller "
         "bottom radius — without incurring additional Euler buckling risk at the thin end.")

    add_figure(doc, "fig_ring_spacing_comparison.png",
               "Figure 4 — Ring spacing comparison: uniform (v3, left) vs geometric L/r "
               "(v4, right).  Geometric spacing keeps L/r constant despite shaft taper, "
               "enabling lighter tapered designs.",
               width_in=5.5)

    para(doc,
         "v4 results: 10 kW winner 10.587 kg (−31.4 % vs v3).  n_lines = 8 selected "
         "unanimously.  target_Lr converged to 2.0 across the winning islands.")

    add_figure(doc, "fig_nlines_mass_curve.png",
               "Figure 5 — Shaft mass vs n_lines for v4 and v5 campaigns overlaid.  "
               "Both campaigns show a clear minimum at n_lines = 8.",
               width_in=5.5)

    # v5
    heading(doc, "3.5  v5 — BEM-coupled rotor radius", level=2)
    para(doc,
         "v4 used a fixed rotor radius R (determined externally from power requirements).  "
         "v5 closes the aerodynamic coupling loop by deriving R self-consistently from "
         "n_lines via a Prandtl tip-loss BEM model:")
    p = doc.add_paragraph()
    r = p.add_run("Cp(n_lines) = (16/27) × [1 − exp(−n_lines/2)] × 0.85")
    r.font.name = "Courier New"
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc,
         "R is then computed from P = Cp × ½ρπR²v³.  This means the optimiser must "
         "account for the aerodynamic cost of each choice of n_lines: more lines → "
         "higher Cp → smaller R → lighter shaft, but also higher tether tension per unit "
         "area.  The self-consistent solution requires slightly larger R than v4 at 10 kW "
         "(5.12 m vs 5.0 m), adding +8.3 % to the 10 kW mass.  At 50 kW, BEM coupling "
         "dramatically improves accuracy and reduces mass by 51 %.")

    add_figure(doc, "fig_cp_contour.png",
               "Figure 6 — BEM Cp as a function of n_lines and TSR.  "
               "The Prandtl tip-loss model shows diminishing returns above n_lines = 8.",
               width_in=5.0)

    add_figure(doc, "fig_bem_analysis.png",
               "Figure 7 — BEM model analysis: rotor radius R and shaft mass vs n_lines "
               "under self-consistent aerodynamic coupling.",
               width_in=5.5)

    doc.add_page_break()

    # ── 4. Structural analysis ─────────────────────────────────────────────────
    heading(doc, "4  Structural Analysis")
    hline(doc)

    heading(doc, "4.1  FOS profile and design space", level=2)
    para(doc,
         "The structural optimiser enforces two simultaneous feasibility constraints: "
         "Euler column buckling (FOS ≥ 1.8) and Tulloch torsional stability (FOS ≥ 1.5).  "
         "Both constraints are evaluated at every ring-to-ring segment.")

    add_figure(doc, "fig_structural_efficiency_profile.png",
               "Figure 8 — FOS profile along the shaft length for the v5 10 kW winner.  "
               "Both Euler and torsional FOS are at or near their respective lower bounds, "
               "indicating a structurally efficient (not over-designed) solution.",
               width_in=5.5)

    add_figure(doc, "fig_fos_landscape.png",
               "Figure 9 — Euler + torsional FOS landscape across the (r_hub, r_bottom) "
               "design space.  The feasible region is bounded by both constraints.",
               width_in=5.5)

    add_figure(doc, "fig_design_space.png",
               "Figure 10 — Feasible region in (r_hub, r_bottom) space.  "
               "The optimum lies in the corner where both constraints are simultaneously "
               "active.",
               width_in=5.0)

    heading(doc, "4.2  Campaign geometry evolution", level=2)
    add_figure(doc, "fig_campaign_geometry_evolution.png",
               "Figure 11 — Winning shaft geometry across v2–v5.  The shift from the "
               "near-cylindrical v3 geometry to the tapered v4/v5 geometry is clearly "
               "visible.",
               width_in=5.5)

    heading(doc, "4.3  Elevation angle trade-off", level=2)
    add_figure(doc, "fig_elevation_angle_trade.png",
               "Figure 12 — Shaft mass and aerodynamic performance vs elevation angle β.  "
               "β = 30° was used throughout v2–v5; the optimum is estimated near 26°.",
               width_in=5.5)

    doc.add_page_break()

    # ── 5. Winning designs ─────────────────────────────────────────────────────
    heading(doc, "5  Winning Designs — v5 Reference")
    hline(doc)

    heading(doc, "5.1  10 kW winner", level=2)
    summary_table(doc,
        ["Parameter", "Value"],
        list(V5_10KW_WINNER.items()),
        col_widths_in=[2.8, 4.0])

    heading(doc, "5.2  50 kW winner", level=2)
    summary_table(doc,
        ["Parameter", "Value"],
        list(V5_50KW_WINNER.items()),
        col_widths_in=[2.8, 4.0])

    doc.add_page_break()

    # ── 6. Open questions ──────────────────────────────────────────────────────
    heading(doc, "6  Open Questions for v6")
    hline(doc)
    para(doc,
         "Three issues must be resolved before committing to hardware.  They are listed "
         "in priority order:")

    for i, (title, body) in enumerate(V6_OPEN_QUESTIONS, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}.  {title}")
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        r.font.color.rgb = NAVY
        r.bold = True
        p.paragraph_format.space_before = Pt(8)
        para(doc, body, size=11)

    doc.add_page_break()

    # ── 7. Appendix — code and data ────────────────────────────────────────────
    heading(doc, "Appendix — Code and Data Reference")
    hline(doc)

    summary_table(doc,
        ["File", "Purpose"],
        [
            ["src/ring_spacing.jl",          "v4/v5 geometric L/r ring spacing; TRPTDesignV4; evaluate_design"],
            ["src/bem.jl",                   "v5 Prandtl tip-loss BEM Cp(n_lines); self-consistent R"],
            ["src/trpt_axial_profiles.jl",   "Torsional collapse constraint (v3+)"],
            ["src/trpt_optimization.jl",     "EvalResult struct; v2/v3 objectives"],
            ["test/test_ring_spacing_v4.jl", "368 unit tests for ring_spacing_v4 (all passing)"],
            ["scripts/run_v4_campaign.jl",   "v4 60-island DE campaign launcher"],
            ["scripts/run_v5_campaign.jl",   "v5 60-island DE campaign launcher"],
            ["scripts/torsional_collapse_check.jl", "Post-hoc v2 torsional validation"],
            ["scripts/results/trpt_opt_v4/", "v4 island results (committed to master)"],
            ["scripts/results/trpt_opt_v5/", "v5 island results (committed to master)"],
            ["DECISIONS.md",                 "Full derivation and rationale for all campaigns"],
        ],
        col_widths_in=[3.0, 4.7])

    doc.add_paragraph()
    para(doc,
         f"Generated {datetime.date.today().isoformat()} by produce_awes_forum_report.py",
         italic=True, color=SLATE, size=9,
         align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()

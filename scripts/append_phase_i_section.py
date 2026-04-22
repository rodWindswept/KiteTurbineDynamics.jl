"""
Append Phase I section to TRPT_Design_Cartography_Report.docx
Covers: Torsional Collapse Constraint (v3 optimisation campaign)

Usage:
  python3 scripts/append_phase_i_section.py
"""

from __future__ import annotations
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import numpy as np

ROOT    = Path(__file__).parent.parent
DOCX    = ROOT / "scripts" / "results" / "trpt_opt_v2" / "TRPT_Design_Cartography_Report.docx"
V2_DIR  = ROOT / "scripts" / "results" / "trpt_opt_v2"
V3_CART = ROOT / "scripts" / "results" / "trpt_opt_v3" / "cartography"

NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
TEAL  = RGBColor(0x00, 0x7A, 0x87)
SLATE = RGBColor(0x44, 0x4F, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MID   = RGBColor(0xD6, 0xE4, 0xED)
GREEN = RGBColor(0x2B, 0x8A, 0x3E)
AMBER = RGBColor(0xE0, 0x7B, 0x00)


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    return p


def para(doc: Document, text: str, bold: bool = False, italic: bool = False,
          color: RGBColor = SLATE, size: int = 11, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    if align is not None:
        p.alignment = align
    return p


def caption(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = SLATE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def load_comparison():
    BEAMS  = ["airfoil", "circular", "elliptical"]
    AXIALS = ["elliptic", "linear", "parabolic", "straight_taper", "trumpet"]
    SEEDS  = ["s1", "s2"]
    SIZES  = ["10kw", "50kw"]

    rows = []
    v3_dir_local = ROOT / "scripts" / "results" / "trpt_opt_v3"
    v3_dir_sibling = ROOT.parent.parent / ".claude" / "worktrees" / "relaxed-shamir-cd9a0f" / "scripts" / "results" / "trpt_opt_v3"
    # Check sibling worktree path
    _check = v3_dir_local / "10kw_circular_straight_taper_s1" / "best_design.json"
    v3_dir = v3_dir_local if _check.exists() else v3_dir_sibling

    for size in SIZES:
        for beam in BEAMS:
            for axial in AXIALS:
                for seed in SEEDS:
                    tag = f"{size}_{beam}_{axial}_{seed}"
                    p2 = V2_DIR / tag / "best_design.json"
                    p3 = v3_dir  / tag / "best_design.json"
                    if not p2.exists() or not p3.exists():
                        continue
                    d2 = json.loads(p2.read_text())
                    d3 = json.loads(p3.read_text())
                    tors = d3["evaluation"].get("torsional_fos_min",
                             d3.get("torsional_fos_min", float("nan")))
                    rows.append({
                        "size": size, "beam": beam, "axial": axial, "seed": seed,
                        "v2_mass": d2["best_mass_kg"],
                        "v3_mass": d3["best_mass_kg"],
                        "v2_taper": d2["design"]["taper_ratio"],
                        "v3_taper": d3["design"]["taper_ratio"],
                        "v2_r_hub": d2["design"]["r_hub_m"],
                        "v3_r_hub": d3["design"]["r_hub_m"],
                        "v3_tors": tors,
                    })
    return rows


def add_summary_table(doc: Document, rows: list[dict]):
    BEAM_L = {"airfoil": "Airfoil", "circular": "Circular", "elliptical": "Elliptical"}
    AX_L   = {"elliptic": "Elliptic", "linear": "Linear", "parabolic": "Parabolic",
               "straight_taper": "Str. Taper", "trumpet": "Trumpet"}

    # One row per size×beam×axial (average over seeds)
    summary = {}
    for r in rows:
        key = (r["size"], r["beam"], r["axial"])
        if key not in summary:
            summary[key] = {"v2": [], "v3": [], "v2t": [], "v3t": [], "v2r": [], "v3r": [], "tors": []}
        summary[key]["v2"].append(r["v2_mass"])
        summary[key]["v3"].append(r["v3_mass"])
        summary[key]["v2t"].append(r["v2_taper"])
        summary[key]["v3t"].append(r["v3_taper"])
        summary[key]["v2r"].append(r["v2_r_hub"])
        summary[key]["v3r"].append(r["v3_r_hub"])
        if not np.isnan(r["v3_tors"]):
            summary[key]["tors"].append(r["v3_tors"])

    headers = ["Config", "Beam", "Axial", "v2 mass (kg)", "v3 mass (kg)",
               "Δ mass %", "v2 taper", "v3 taper", "v3 Tor FOS"]
    table = doc.add_table(rows=1 + len(summary), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, ((size, beam, axial), vals) in enumerate(sorted(summary.items())):
        m2   = np.mean(vals["v2"])
        m3   = np.mean(vals["v3"])
        pct  = 100 * (m3 - m2) / m2
        t2   = np.mean(vals["v2t"])
        t3   = np.mean(vals["v3t"])
        tors = np.mean(vals["tors"]) if vals["tors"] else float("nan")

        bg = MID if i % 2 == 0 else WHITE
        data = [size.upper(), BEAM_L[beam], AX_L[axial],
                f"{m2:.2f}", f"{m3:.2f}", f"+{pct:.0f}%",
                f"{t2:.3f}", f"{t3:.3f}",
                f"{tors:.3f}" if not np.isnan(tors) else "—"]

        row_cells = table.rows[i+1].cells
        for j, val in enumerate(data):
            set_cell_bg(row_cells[j], bg)
            p = row_cells[j].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.font.name = "Calibri"
            run.font.color.rgb = SLATE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Highlight % change in amber
            if j == 5:
                run.font.color.rgb = AMBER
                run.font.bold = True

    return table


def main():
    doc = Document(DOCX)

    heading(doc, "Phase I — Torsional Collapse Constraint (v3)", level=1)

    para(doc,
         "Background: what was missing in v2",
         bold=True, color=NAVY, size=11)

    para(doc, (
        "All v2 optimisations (Phases C–H) minimised total beam mass subject to a "
        "bending/buckling factor of safety FOS ≥ 1.8 at every ring.  Torsional collapse "
        "was not constrained.  Torsional collapse occurs when the shaft torque exceeds the "
        "critical torque that causes a cylindrical (or conical) shell to buckle under "
        "combined axial compression and torsion.  Without this constraint the optimiser "
        "exploited highly conical (low taper_ratio ≈ 0.25) geometries, which are "
        "dramatically lighter under bending alone but fail torsionally at a fraction of "
        "design load.  54 of the 60 v2 winning designs had taper_ratio < 0.9 and would "
        "collapse torsionally in service."
    ), color=SLATE)

    para(doc,
         "Tulloch torsional-collapse criterion",
         bold=True, color=NAVY, size=11)

    para(doc, (
        "v3 enforces the Tulloch criterion: the torsional factor of safety "
        "τ_FOS = τ_critical / τ_applied ≥ 1.5 at every ring.  τ_critical is the "
        "Tulloch buckling torque for a thin-walled tube under combined bending and "
        "torsion; τ_applied is the TRPT shaft torque at that ring cross-section.  "
        "The constraint is evaluated inside the fitness function alongside the existing "
        "FOS ≥ 1.8 bending check; any design violating either constraint is marked "
        "infeasible and discarded by the optimiser."
    ), color=SLATE)

    para(doc,
         "What changed: taper forced to cylindrical",
         bold=True, color=NAVY, size=11)

    para(doc, (
        "The torsional constraint is catastrophic for conical shafts.  A conical TRPT "
        "shaft tapers from a large ground-ring radius to a small hub radius; its thin "
        "top rings see the same shaft torque as the bottom rings but have far smaller "
        "cross-section, making them the governing failure mode.  The only way the "
        "optimiser can satisfy τ_FOS ≥ 1.5 across all rings is to make the shaft "
        "cylindrical (taper_ratio = 1.0) — confirmed by all 60 v3 designs converging "
        "to taper_ratio = 1.000."
    ), color=SLATE)

    para(doc,
         "Quantitative mass impact",
         bold=True, color=NAVY, size=11)

    para(doc, (
        "Table I-1 summarises v2 vs v3 masses for all 15 beam×axial combinations at "
        "10 kW and 50 kW.  The 10 kW circular straight-taper winner increased from "
        "2.81 kg (v2, taper=0.25) to 15.44 kg (v3, taper=1.0) — a +449 % increase. "
        "The mean mass increase across all 60 configurations is +304 %.  Six 50 kW "
        "airfoil configurations already had taper_ratio = 1.0 in v2 and are unaffected "
        "(0 % change).  All others are significantly heavier under the physical constraint."
    ), color=SLATE)

    doc.add_paragraph()  # space
    add_summary_table(doc, load_comparison())
    caption(doc, "Table I-1.  v2 vs v3 mass and geometry comparison for all 60 configurations "
                 "(values averaged over two random seeds; taper column shows mean).")
    doc.add_paragraph()

    para(doc,
         "The taper_ratio = 1.0 finding and its physical meaning",
         bold=True, color=NAVY, size=11)

    para(doc, (
        "A cylindrical TRPT shaft has constant cross-sectional moment capacity from "
        "ground to hub.  This means the shaft wall, once sized to resist the peak "
        "torque at the ground ring, is over-designed at every higher ring — the shaft "
        "carries excess material throughout its length.  This is a fundamental mass "
        "penalty imposed by the torsional physics.  Future optimisation must explore "
        "alternative structural strategies (variable wall thickness at constant outer "
        "diameter, composite layup, lattice rings) to recover some of the mass advantage "
        "of v2 without violating torsional collapse."
    ), color=SLATE)

    para(doc,
         "Practical implication for ground ring deployment",
         bold=True, color=NAVY, size=11)

    para(doc, (
        "A cylindrical shaft means that all rings have the same radius.  This "
        "significantly simplifies ground ring deployment: the ground ring (the lowest "
        "ring, previously the widest in a conical design) is now the same diameter as "
        "all upper rings.  Ground logistics (transportation, anchoring radius, "
        "swept-area footprint) are unchanged relative to the hub radius rather than "
        "expanding to a wide base.  The torsional constraint thus has a practical "
        "benefit for site deployment, even though it increases mass."
    ), color=SLATE)

    # Figures
    doc.add_paragraph()
    para(doc, "Figures — Phase I", bold=True, color=NAVY, size=11)

    figs = [
        ("fig_v2_vs_v3_mass_comparison.png",
         "Figure I-1.  Mass impact of enforcing the Tulloch torsional collapse "
         "constraint.  Grey bars: v2 (no torsional constraint).  Teal bars: v3 "
         "(τ_FOS ≥ 1.5).  Percentage labels show mean mass increase per combination. "
         "The 50 kW airfoil elliptic/linear/trumpet bars are equal height because "
         "those v2 designs already had taper=1.0."),
        ("fig_v3_geometry_shift.png",
         "Figure I-2.  Geometry change driven by the torsional constraint.  "
         "Left: taper_ratio scatter — all 60 v3 designs sit at taper=1.0 (red "
         "dotted line) regardless of v2 taper, confirming the constraint forces "
         "cylindrical geometry.  Right: hub radius shift — v3 r_hub is consistently "
         "larger than v2 r_hub as the optimiser compensates for the cylindrical "
         "constraint by growing the shaft radius."),
        ("fig_v3_winner_glmakie.png",
         "Figure I-3.  GLMakie render of the v3 10 kW winner: circular straight-taper, "
         "15.44 kg, FOS=1.8, torsional FOS=1.5, 5 rings × 8 lines, r_hub=1.99 m, "
         "taper=1.0 (cylindrical).  Beam outer diameter colour-coded (viridis); "
         "knuckle vertices as red spheres; tethers as dark grey lines."),
    ]

    for fname, cap_text in figs:
        img_path = V3_CART / fname
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(5.5))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para(doc, f"[Image not found: {fname}]", italic=True, color=AMBER)
        caption(doc, cap_text)
        doc.add_paragraph()

    doc.save(DOCX)
    print(f"Saved Phase I section → {DOCX}")


if __name__ == "__main__":
    main()

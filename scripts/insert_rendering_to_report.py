# scripts/insert_rendering_to_report.py
import sys
from docx import Document
from docx.shared import Inches

def insert_render():
    print("------------------------------------------")
    print("Inserting GLMakie render into Lift_Kite_Sizing_Report.docx...")
    print("------------------------------------------")
    
    doc = Document("Lift_Kite_Sizing_Report.docx")
    
    target_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if "A Julia GLMakie render of the installed geometry" in p.text:
            target_idx = idx
            break
            
    if target_idx == -1:
        print("Error: Could not find the placeholder paragraph in Lift_Kite_Sizing_Report.docx!")
        return
        
    print(f"Found placeholder paragraph at index {target_idx}:")
    print(f"  '{doc.paragraphs[target_idx].text}'")
    
    # 1. Update the placeholder text to refer to Figure 6
    p = doc.paragraphs[target_idx]
    p.text = p.text.replace(
        "A Julia GLMakie render of the installed geometry (with annotated dimensions) will be added to this section in a subsequent revision.",
        "The Julia GLMakie 3D system rendering of the installed geometry is shown below in Figure 6."
    )
    
    # 2. Insert image paragraph and caption paragraph right after the text
    # We do this by inserting before paragraph (target_idx + 1)
    next_p = doc.paragraphs[target_idx + 1]
    
    # Insert image paragraph
    p_image = next_p.insert_paragraph_before()
    p_image.alignment = 1  # 1 = Center
    run = p_image.add_run()
    run.add_picture("figures/fig_trpt_installed_geometry.png", width=Inches(5.8))
    
    # Insert caption paragraph
    p_caption = next_p.insert_paragraph_before()
    p_caption.alignment = 1  # Center
    run_cap = p_caption.add_run("Figure 6: Julia GLMakie 3D system rendering of the installed 10 kW TRPT autogyro kite turbine system showing ground anchoring, the multi-ring TRPT transmission, and the rotary autogyro lifter.")
    run_cap.font.italic = True
    run_cap.font.size = Inches(0.12)  # subtle, clean caption font size (~9 pt)
    
    # Add a blank paragraph for spacing
    p_space = next_p.insert_paragraph_before()
    p_space.text = ""
    
    doc.save("Lift_Kite_Sizing_Report.docx")
    print("Lift_Kite_Sizing_Report.docx successfully updated with GLMakie 3D rendering!")

if __name__ == "__main__":
    insert_render()

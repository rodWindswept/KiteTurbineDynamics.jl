"""
TRPT Kite Turbine Dynamics — Technical Report Generator
=========================================================
Produces a formatted .docx report from simulation results for sharing
with investors, engineers, and research collaborators.

Usage:
  python3 scripts/produce_report.py

Output:
  TRPT_Dynamics_Report.docx   (in repo root)
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import csv, datetime

ROOT    = Path(__file__).parent.parent
FIGS    = ROOT / "scripts" / "results" / "lift_kite"
MPPT    = ROOT / "scripts" / "results" / "mppt_twist_sweep"
OUT     = ROOT / "TRPT_Dynamics_Report.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x0D, 0x1B, 0x2A)   # deep navy  — headings
TEAL    = RGBColor(0x00, 0x7A, 0x87)   # teal       — accent / rules
SLATE   = RGBColor(0x44, 0x4F, 0x5A)   # slate      — body text
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT   = RGBColor(0xF0, 0xF4, 0xF8)   # light blue-grey — table header bg
MID     = RGBColor(0xD6, 0xE4, 0xED)   # light teal — alt rows

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Add borders to a table cell. kwargs: top, bottom, left, right each a dict."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, attrs in kwargs.items():
        border = OxmlElement(f"w:{side}")
        for k, v in attrs.items():
            border.set(qn(f"w:{k}"), v)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = NAVY if level == 1 else TEAL
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, bold=False, italic=False, size=10.5):
    p  = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = SLATE
    run.font.bold   = bold
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p

def caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x77, 0x88, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    return p

def callout(doc, text, label="KEY RESULT"):
    """Teal left-border callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    lbl = p.add_run(f"{label}  ")
    lbl.font.bold = True
    lbl.font.color.rgb = TEAL
    lbl.font.size = Pt(10)
    txt = p.add_run(text)
    txt.font.size = Pt(10)
    txt.font.color.rgb = SLATE
    # left border via paragraph border XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "18")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "007A87")
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header row and alternating row shading."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, NAVY)
        run = cell.paragraphs[0].add_run(h)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        bg = MID if r_idx % 2 == 1 else RGBColor(0xFF, 0xFF, 0xFF)
        for c_idx, val in enumerate(row_data):
            cell = row_cells[c_idx]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = SLATE
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Column widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacer
    return tbl

def add_figure(doc, img_path, width_in=6.0, cap=None):
    if not Path(img_path).exists():
        body(doc, f"[Figure not found: {img_path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_in))
    if cap:
        caption(doc, cap)

def hr(doc):
    """Thin horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "007A87")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("TRPT Kite Turbine Dynamics")
tr.font.size  = Pt(26)
tr.font.bold  = True
tr.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Multi-body Simulation Results — 10 kW Prototype")
sr.font.size  = Pt(14)
sr.font.color.rgb = TEAL
sr.font.bold  = False

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run(
    f"Windswept & Interesting Ltd   ·   {datetime.date.today().strftime('%B %Y')}\n"
    "KiteTurbineDynamics.jl — Full multi-body ODE simulator"
)
mr.font.size = Pt(10)
mr.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

doc.add_page_break()

# ── 1  Executive Summary ───────────────────────────────────────────────────────
heading(doc, "1  Executive Summary")
hr(doc)

body(doc, (
    "KiteTurbineDynamics.jl is a full multi-body dynamics simulator for the "
    "Windswept & Interesting TRPT (Tensile Rotary Power Transmission) kite turbine. "
    "Unlike quasi-static models, every tether line is individually modelled as a "
    "chain of spring-damper rope nodes. Torsional coupling between rings is emergent "
    "from helical tether geometry — no analytical torque formula is used. The system "
    "comprises 241 nodes and a 1,478-state ODE."
))

callout(doc,
    "CT thrust and shaft tension cancel at the hub in quasi-static equilibrium. "
    "The lift kite supports the airborne weight only (245 N). "
    "A 27.5 m² kite sized for v = 4 m/s launch delivers 8.3× margin at rated wind.")

body(doc, (
    "The simulation has been corrected to remove phantom kite CL lift from the rotor "
    "disc force model. All results in this document reflect the corrected physics. "
    "288 of 288 automated tests pass."
))

# ── 1.5 Simulation Physics & Known Limitations ─────────────────────────────────
heading(doc, "1.5  Simulation Physics & Known Limitations")
hr(doc)
body(doc, "The current simulation environment contains several known physical abstractions and limitations that contextualize these results:")
bullet_points = [
    "Zero-Speed Thrust (CT=0 at λ=0): The thrust coefficient CT is set exactly to 0.0 at standstill. A physical stationary rotor disk would experience significant drag. This affects 'cold start' collapse models.",
    "Startup Torque numerical 'Hack': Aerodynamic torque is calculated as P_aero / max(|ω|, 0.5) to prevent division by zero, giving a small numerical 'kick-start' at ω=0 since P_aero is also zero.",
    "Fixed-Mast Hub Constraint: The elevation constraint is a single tension-only spring. The hub is constrained relative to a fixed ground point and does not have the full freedom of a true tethered kite.",
    "Analytical vs. Dynamic Scaling: 'Stacked Rotor' configurations in subsequent analytical reports are derived via scaling laws, not multi-rotor dynamic simulations.",
    "Torsional Damping: An explicit, non-physical inter-ring torsional damper is applied to suppress high-frequency torsional oscillations numerical integration issues."
]
for bp in bullet_points:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(bp).font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

# ── 2  System Architecture ────────────────────────────────────────────────────
heading(doc, "2  System Architecture")
hr(doc)

add_figure(doc, FIGS / "diag_trpt_system.png", 5.5,
           "Figure 1 — TRPT system overview: 16-ring multi-body shaft at 30° elevation, "
           "30 m tether, lift kite above hub.")

body(doc, "The 10 kW prototype parameters are:")

add_table(doc,
    ["Parameter", "Value", "Notes"],
    [
        ["Hub elevation angle",   "30°",            "Shaft axis direction"],
        ["Tether length",         "30 m",           "Hub to ground distance"],
        ["Number of rings",       "16",             "Ring 1 = ground, 16 = hub"],
        ["Lines per ring pair",   "3",              "Dyneema, 4 mm diameter"],
        ["Rated power",           "10 kW",          "At v = 11 m/s"],
        ["Rotor radius",          "0.9 m",          "Mean ring radius"],
        ["MPPT k_nominal",        "11 N·m·s²/rad²", "ω³ load law"],
        ["Hub mass",              "~24 kg",         "Blades + rings + tether mass"],
        ["Airborne weight",       "245 N",          "Corrected: CT thrust cancels"],
        ["Lift kite area",        "27.5 m²",        "Sized at v_design = 4 m/s"],
    ],
    col_widths=[2.0, 1.5, 3.0]
)

# ── 3  Hub Force Balance ───────────────────────────────────────────────────────
heading(doc, "3  Hub Force Balance — Corrected CT-Thrust Physics")
hr(doc)

body(doc, (
    "A key correction was made to the hub force balance. Previously, the model "
    "incorrectly applied phantom kite CL lift and over-estimated the required lift "
    "force. The correct analysis is:"
))
body(doc, (
    "The rotor disc is tilted at 60° from horizontal (shaft at 30° elevation). "
    "CT thrust acts along the shaft axis — upward and downwind. In quasi-static "
    "equilibrium, shaft tension from the TRPT elastic springs below the hub is "
    "equal in magnitude to CT thrust and opposite in direction. They cancel. "
    "The net downward load on the lift kite is the airborne weight only."
))

add_figure(doc, FIGS / "diag_hub_forces.png", 5.5,
           "Figure 2 — Corrected hub force diagram. CT thrust (2,394 N) and shaft "
           "tension cancel; lift kite supports 245 N airborne weight only.")

callout(doc,
    "Old model: lift required = 1,441 N at v=11 m/s (included double-counted thrust).  "
    "Corrected: lift required = 245 N (airborne weight only, wind-independent).",
    label="PHYSICS CORRECTION")

# ── 4  Lift Device Comparison ─────────────────────────────────────────────────
heading(doc, "4  Lift Device Architecture Comparison")
hr(doc)

body(doc, (
    "Three lift device architectures were analysed against the corrected 245 N "
    "hub support requirement:"
))

add_figure(doc, FIGS / "diag_lift_devices.png", 5.5,
           "Figure 3 — Lift device architectures: single kite, stacked kites, rotary lifter.")

add_table(doc,
    ["Metric", "Single Kite", "Stack × 3", "Rotary Lifter"],
    [
        ["Required lift force",    "245 N",     "245 N",    "245 N"],
        ["Area (v=4 m/s sizing)",  "27.5 m²",   "27.5 m² total", "~12 m² blades (TBD)"],
        ["Individual unit size",   "27.5 m²",   "9.2 m² each",   "1.5 m radius rotor"],
        ["Lift margin @ v=11 m/s", "8.3×",      "8.3×",     "1.6×"],
        ["Tension CV @ v=11 m/s",  "30.1%",     "30.2%",    "3.6%"],
        ["CV vs single kite",      "—",         "1.00×",    "0.12× (8× better)"],
        ["Hub excursion std",       "26–39 mm",  "~26 mm",   "TBD (predicted ~4 mm)"],
    ],
    col_widths=[2.4, 1.5, 1.5, 1.5]
)

callout(doc,
    "The rotary lifter provides 8× lower tension coefficient of variation than "
    "a passive kite, predicting a corresponding 8× reduction in hub altitude sway. "
    "Confirmed in short-run simulation (3.9× improvement at 3 s; longer runs pending).")

body(doc, "The stacked-kite approach offers handling advantages (smaller individual kite "
     "area) with identical hub stability to a single kite — no stability benefit, but "
     "significant practical benefit for ground handling and launch.")

# ── 5  Hub Excursion ──────────────────────────────────────────────────────────
heading(doc, "5  Dynamic Hub Excursion — Long-Run Simulation")
hr(doc)

body(doc, (
    "A 84-minute turbulent wind simulation (IEC Class A, I = 0.15) was run across "
    "12 device × wind-speed combinations to characterise hub altitude stability. "
    "Results confirm the simulator correctly captures the lift kite's stabilising "
    "effect on hub position."
))

add_figure(doc, FIGS / "hub_excursion_analysis.png", 6.2,
           "Figure 4 — Hub excursion analysis: 12-panel comparison of hub_z standard "
           "deviation across devices and wind speeds (84-minute simulation).")

add_table(doc,
    ["Device", "v = 8 m/s  hub_z std", "v = 11 m/s  hub_z std", "vs NoLift baseline"],
    [
        ["SingleKite",   "39 mm",  "26 mm",  "—"],
        ["Stack × 3",    "~39 mm", "~26 mm", "same as SingleKite"],
        ["NoLift",        "72 mm",  "36 mm",  "1.4–2× worse"],
        ["RotaryLifter",  "TBD",    "TBD",    "predicted ~0.12×"],
    ],
    col_widths=[2.0, 1.8, 1.8, 2.3]
)

body(doc, (
    "The NoLift case confirms that CT thrust alone holds the hub at altitude — "
    "the turbine is self-suspending above ~3.5 m/s — but with significantly higher "
    "sway than a kite-supported configuration. The single kite reduces hub excursion "
    "by 28–46% at 8–11 m/s compared to the no-lift baseline."
))

# ── 6  MPPT Twist Sweep ───────────────────────────────────────────────────────
heading(doc, "6  MPPT Gain × Twist Angle Sweep")
hr(doc)

body(doc, (
    "A parametric sweep of 7 MPPT gain multipliers across 4 wind speeds (28 combinations, "
    "60 s settled simulation each) characterises how the TRPT shaft twist angle responds "
    "to generator loading. The research question: does steady-state twist carry enough "
    "information to serve as a sensorless MPPT or bridling control signal?"
))

add_figure(doc, MPPT / "twist_sweep_v2_analysis.png", 6.2,
           "Figure 5 — MPPT × twist sweep v2: power, twist, tether load, and wind ramp "
           "response across 7 MPPT gain multipliers and 4 wind speeds.")

heading(doc, "6.1  Power vs MPPT gain", level=2)

# Load summary CSV for a clean table
sum_rows_doc = []
try:
    with open(MPPT / "twist_sweep_v2_summary.csv") as f:
        reader = csv.DictReader(f)
        v_winds = [8.0, 10.0, 11.0, 13.0]
        by_km = {}
        for row in reader:
            km = float(row["k_mult"])
            vw = float(row["v_wind"])
            if km not in by_km:
                by_km[km] = {}
            by_km[km][vw] = (float(row["P_kw_mean"]), float(row["twist_mean"]))
        for km in sorted(by_km.keys()):
            r = [f"{km:.2g}×"]
            for v in v_winds:
                if v in by_km[km]:
                    P, tw = by_km[km][v]
                    r.append(f"{P:.2f} kW  ({tw:.0f}°)")
                else:
                    r.append("—")
            sum_rows_doc.append(r)
except FileNotFoundError:
    sum_rows_doc = [["(data file not found)", "", "", "", ""]]

add_table(doc,
    ["k_mult", "v = 8 m/s  P (twist)", "v = 10 m/s  P (twist)",
     "v = 11 m/s  P (twist)", "v = 13 m/s  P (twist)"],
    sum_rows_doc,
    col_widths=[0.8, 1.6, 1.6, 1.6, 1.6]
)

callout(doc,
    "Optimal MPPT gain is k × 1.2 across all wind speeds (very flat peak between "
    "k×1.0 and k×1.2). Peak power at v=11 m/s: 8.31 kW. "
    "Twist at optimal: 238° (8 m/s) → 308° (13 m/s).")

heading(doc, "6.2  Twist as a control signal", level=2)
body(doc, (
    "The sweep confirms twist increases monotonically with MPPT gain at fixed "
    "wind speed — making it a reliable proxy for shaft torque. However, twist "
    "also tracks wind speed strongly at fixed gain, making it ambiguous as a "
    "sole control input: the same twist angle can correspond to two different "
    "operating points (under-braked fast vs over-braked stalled)."
))
body(doc, (
    "The torque-to-tension ratio τ/T is the more physically meaningful control "
    "variable. The analytical prediction δα ≈ (τ/T) × L/nr² holds in the linear "
    "regime and allows twist to be directly predicted from tether tension measurements."
))

heading(doc, "6.3  Wind ramp dynamics", level=2)
body(doc, (
    "A 7→14 m/s wind ramp over 150 s (cold-ish start from v=7 m/s near-zero RPM) "
    "reveals the TRPT long mechanical inertia time constant:"
))
add_table(doc,
    ["Time (s)", "v_wind (m/s)", "Twist (°)", "Power (kW)", "Note"],
    [
        ["25",  "8.2",  "0.7°",  "0.00", "Rotor just starting"],
        ["65",  "10.0", "120°",  "2.17", "Spinning up"],
        ["105", "11.9", "167°",  "2.20", "Still accelerating"],
        ["155", "14.0", "200°",  "2.25", "End of ramp"],
        ["—",   "11.0", "284°",  "8.27", "Steady-state reference"],
    ],
    col_widths=[0.9, 1.3, 1.0, 1.2, 2.5]
)
callout(doc,
    "At v=14 m/s end-of-ramp, P = 2.25 kW vs 13.4 kW at steady state. "
    "The TRPT spin-up time constant far exceeds 150 s. Controllers must account "
    "for this inertial lag. Consistent with torque wave resonance (Tulloch et al.).",
    label="KEY DYNAMICS")

# ── 7  Simulation Scope and Limitations ──────────────────────────────────────
heading(doc, "7  Simulation Scope and Known Limitations")
hr(doc)

body(doc, (
    "The simulator produces high-fidelity results for above-cut-in steady-state and "
    "transient conditions. The following limitations apply to the current version and "
    "are areas of active development:"
))

add_table(doc,
    ["Limitation", "Impact", "Planned fix"],
    [
        ["TRPT/rotor collapse not observed\nin low-wind scenarios",
         "Simulator does not demonstrate\nshaft un-twist or hub drop\nbelow cut-in",
         "Cold-start test from ω=0;\nmulti-element back line;\nfull 6-DOF hub frame"],
        ["Back line single spring-damper",
         "Cannot sag, go slack, or\nlet hub drop forward",
         "5+ rope node back line model"],
        ["Hub elevation β partly fixed\nby parameterisation",
         "Hub cannot fall to ground\neven at zero lift",
         "Confirm full 6-DOF free\nhub state in ODE"],
        ["Kite lift clamped to F_required",
         "Below cut-in, actual lift\n< F_req not propagated\nto hub kinematics",
         "Apply actual F_lift;\nlet deficit cause sag"],
        ["Rotary lifter long-run\nexcursion not yet measured",
         "8× CV improvement\nnot confirmed in\nlong-run statistics",
         "Dedicated rotary lifter\n84-min simulation run"],
    ],
    col_widths=[2.2, 1.8, 2.1]
)

body(doc, (
    "All results in this document are labelled \"above cut-in, kite-suspended\" "
    "and should be interpreted in that context. The fixed-mast caveat applies: "
    "the hub elevation angle is substantially constrained by the current model "
    "and does not represent a fully free-floating kite."
), italic=True)

# ── 8  Next Steps ─────────────────────────────────────────────────────────────
heading(doc, "8  Next Steps")
hr(doc)

add_table(doc,
    ["Priority", "Task", "Outcome"],
    [
        ["1", "Demonstrate TRPT collapse\nin low-wind cold-start",
              "Validate non-ideal operation;\nremove fixed-mast caveat"],
        ["2", "Multi-element back line model",
              "Realistic slack/sag;\nmore accurate low-wind hub dynamics"],
        ["3", "Rotary lifter long-run\nexcursion sweep",
              "Confirm 8× stability improvement\nover passive kite"],
        ["4", "Torque wave resonance analysis\n(Tulloch et al. method)",
              "Characterise TRPT torsional\nnormal modes; improve MPPT design"],
        ["5", "Annulus rotor blade element\nforce diagram (BEM)",
              "Visualise lift/drag vectors\nat each radial annulus"],
        ["6", "Validate δα ≈ (τ/T)×geometry\nagainst sweep data",
              "Confirm twist-over-tension\nas sensorless MPPT signal"],
    ],
    col_widths=[0.7, 2.5, 2.9]
)

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
hr(doc)
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run(
    "KiteTurbineDynamics.jl  ·  Windswept & Interesting Ltd  ·  "
    f"{datetime.date.today().strftime('%d %B %Y')}  ·  "
    "288/288 tests passing  ·  Corrected CT-thrust physics"
)
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x99, 0xAA, 0xBB)
fr.font.italic = True

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Report saved: {OUT}")

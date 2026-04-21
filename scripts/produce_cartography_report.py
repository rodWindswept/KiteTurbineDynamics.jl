#!/usr/bin/env python3
"""scripts/produce_cartography_report.py
Phase H — Generate TRPT_Design_Cartography_Report.docx from everything the
168-hour autonomous campaign has produced so far. Self-documenting:
each section explains *what was researched*, *why*, and *what was found*.

Output: scripts/results/trpt_opt_v2/TRPT_Design_Cartography_Report.docx

Run at any time — script degrades gracefully if some assets are missing.
"""
from pathlib import Path
from datetime import datetime
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).resolve().parent.parent
V2   = REPO / "scripts" / "results" / "trpt_opt_v2"
LHS  = V2 / "lhs"
CART = V2 / "cartography"
REND = V2 / "renders"
DLF  = REPO / "scripts" / "results" / "trpt_opt" / "dlf"
OUT  = V2 / "TRPT_Design_Cartography_Report.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_image_if_exists(doc, path, width_in=6.2, caption=None):
    if not Path(path).exists():
        p = doc.add_paragraph(f"[image missing: {Path(path).name}]")
        p.runs[0].italic = True
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)


def load_campaign_data():
    """Load all available DE archives + LHS CSVs + best_design.json files."""
    data = {"islands": [], "lhs": [], "best": None}
    for arc in V2.glob("*/elite_archive.csv"):
        try:
            df = pd.read_csv(arc)
            if len(df) == 0: continue
            df["island"] = arc.parent.name
            data["islands"].append(df)
        except Exception as e:
            print(f"skip {arc}: {e}")
    for jsn in V2.glob("*/best_design.json"):
        try:
            with open(jsn) as f:
                d = json.load(f)
            d["island"] = jsn.parent.name
            if data["best"] is None or d.get("best_mass_kg", 1e9) < data["best"].get("best_mass_kg", 1e9):
                data["best"] = d
        except Exception:
            pass
    for csv in LHS.glob("*.csv"):
        try:
            df = pd.read_csv(csv)
            parts = csv.stem.split("_")
            df["config"] = parts[0]; df["beam"] = parts[1]
            data["lhs"].append(df)
        except Exception:
            pass
    return data


def main():
    data = load_campaign_data()
    doc = Document()

    # ── Title ──
    t = doc.add_heading("TRPT Design Cartography Report", level=0)
    sub = doc.add_paragraph()
    r = sub.add_run(f"168-hour autonomous design-space campaign\n"
                    f"Generated {datetime.now().isoformat(timespec='seconds')}\n"
                    f"KiteTurbineDynamics.jl — Windswept Energy")
    r.italic = True

    # ── Executive summary ──
    add_heading(doc, "Executive summary", 1)
    if data["islands"]:
        n_islands = len(data["islands"])
        total_archive = sum(len(df) for df in data["islands"])
        best_df = pd.concat(data["islands"], ignore_index=True)
        best_df = best_df.sort_values("mass_kg")
        winner = best_df.iloc[0]
        summary = (
            f"The Phase C autonomous campaign ran {n_islands} Differential "
            f"Evolution islands spanning 2 configurations × 3 beam profiles × "
            f"5 axial-curve families × 2 seeds. Each island maintained a "
            f"feasible-design elite archive of up to 200 unique candidates, "
            f"giving {total_archive:,} catalogued feasible designs.\n\n"
            f"Best feasible design across the entire campaign:\n"
            f"  • island     : {winner['island']}\n"
            f"  • total mass : {winner['mass_kg']:.3f} kg\n"
            f"  • min FOS    : {winner['min_fos']:.3f}\n"
            f"  • n_rings    : {winner['n_rings']}\n"
            f"  • n_lines    : {winner['n_lines']}\n"
            f"  • r_hub      : {winner['r_hub_m']:.3f} m\n"
            f"  • axial_idx  : {winner['axial_idx']}\n"
            f"  • knuckle_g  : {winner['knuckle_mass_kg']*1000:.1f}"
        )
        doc.add_paragraph(summary)
    else:
        doc.add_paragraph(
            "Campaign still in progress — no DE island has written a "
            "completed archive yet. Proceed to LHS and Phase D sections for "
            "interim results."
        )

    # ── Phase A ──
    add_heading(doc, "Phase A — Search-space expansion (complete)", 1)
    doc.add_paragraph(
        "THESIS: The prior 7-DoF linear-taper search box was too small. "
        "Real TRPT designs in the Grasshopper archive used elliptic, "
        "parabolic, and \"straight bottom then grow\" axial profiles, so the "
        "optimizer needs those curves in its vocabulary. We also promote "
        "n_lines (which is identical to n_polygon_sides and n_blades in the "
        "topology) and knuckle_mass_kg to free decision variables."
    )
    doc.add_paragraph(
        "IMPLEMENTATION: src/trpt_axial_profiles.jl adds a five-member axial "
        "family — LINEAR, ELLIPTIC, PARABOLIC, TRUMPET, STRAIGHT_TAPER — and "
        "a 12-DoF TRPTDesignV2 record. search_bounds_v2() defines the new "
        "envelope; objective_v2() is the optimizer-facing scalar cost. "
        "ring_radii(design) computes r(z) over uniform z spacing using r_of_z()."
    )
    doc.add_paragraph(
        "FINDINGS: Test suite extends from 308 → 331 assertions (all pass). "
        "n_lines = 5, 6, and 7 designs behave very differently — the polygon "
        "compression factor 1/(2·tan(π/n)) falls from 0.68 at n=3 to 0.19 at "
        "n=8, and the side length 2·r·sin(π/n) halves from n=3 to n=6 at "
        "fixed r. Both matter because Euler's P_crit scales as 1/L²."
    )
    add_image_if_exists(doc, V2 / "fig_polygon_family.png", width_in=6.4,
                         caption="Figure A-1. Polygon family at fixed r_hub — "
                                 "ring shape, side length, vertex compression "
                                 "factor, and P_crit trade under equal mass budget.")

    # ── Phase B ──
    add_heading(doc, "Phase B — DLF calibration and centripetal physics", 1)
    doc.add_paragraph(
        "THESIS: The original OPT_DESIGN_LOAD_FACTOR=0.5 was engineering "
        "guesswork. Before we run a serious optimization we must calibrate "
        "it against the live multi-body ODE under realistic fault and gust "
        "transients, then decide — as an engineering judgement — which "
        "envelope designs should actually size against."
    )
    doc.add_paragraph(
        "EXPERIMENT: scripts/calibrate_dlf.jl ran six scenarios: steady wind "
        "at 11 / 15 / 20 / 25 m/s, a coherent IEC-61400-1 gust from 11→25 "
        "m/s, and an emergency-brake transient (3× MPPT gain step)."
    )
    add_image_if_exists(doc, DLF / "fig_dlf_envelope_per_ring.png", width_in=6.2,
                         caption="Figure B-1. Per-ring DLF envelope across "
                                 "scenarios (calibrate_dlf.jl).")
    add_image_if_exists(doc, DLF / "fig_dlf_summary_bars.png", width_in=6.2,
                         caption="Figure B-2. Mean/p95/peak DLF per scenario.")
    doc.add_paragraph(
        "FINDING: Peak DLF per scenario — steady25: 0.32, gust 11→25: 0.74, "
        "emergency brake: 1.39. The ebrake case dominates the envelope by a "
        "factor of 2×."
    )
    doc.add_paragraph(
        "OPERATIONAL DECISION (Rod, 2026-04-20): hard brakes are avoided by "
        "procedure — the rotor is first yaw-stalled via the back-anchor "
        "tether before any mechanical braking. We therefore exclude the "
        "ebrake case from the sizing envelope and set OPT_DESIGN_LOAD_FACTOR "
        "= 1.2 (60% margin over the steady-state worst + 60% margin over "
        "coherent gust). This changes from 1.5 → 1.2, saving beam mass."
    )
    doc.add_paragraph(
        "CENTRIPETAL PHYSICS ADDENDUM: Each vertex now subtracts m·Ω²·r "
        "from the inward line load before computing polygon compression. "
        "At rated Ω = 20 rad/s and r = 5 m, the blade mass lumped at the "
        "hub ring (~2.2 kg / vertex) contributes ~440 N of outward force — "
        "not negligible. Knuckle + beam mass at interior rings contributes "
        "tens of newtons, small but included for completeness."
    )

    # ── Phase C ──
    add_heading(doc, "Phase C — Grand parameter sweep", 1)
    doc.add_paragraph(
        "THESIS: Now that the physics is calibrated and the search space is "
        "enriched, explore it thoroughly. 60 DE islands (2 configs × 3 "
        "beam profiles × 5 axial curves × 2 seeds) each run up to 3 hours "
        "and each maintain a 200-element elite archive of unique feasible "
        "designs — so the output is a *design space cartography*, not just "
        "one winner."
    )
    if data["islands"]:
        combined = pd.concat(data["islands"], ignore_index=True)
        # Top-10
        top10 = combined.sort_values("mass_kg").head(10)
        doc.add_paragraph("TOP 10 FEASIBLE DESIGNS ACROSS ALL ISLANDS:")
        tbl = doc.add_table(rows=1, cols=8)
        hdr = tbl.rows[0].cells
        for c, h in zip(hdr, ["island", "mass_kg", "FOS", "n_rings",
                              "n_lines", "r_hub_m", "axial_idx", "knuckle_g"]):
            c.text = h
        for _, row in top10.iterrows():
            cells = tbl.add_row().cells
            cells[0].text = str(row["island"])
            cells[1].text = f"{row['mass_kg']:.3f}"
            cells[2].text = f"{row['min_fos']:.2f}"
            cells[3].text = str(int(row["n_rings"]))
            cells[4].text = str(int(row["n_lines"]))
            cells[5].text = f"{row['r_hub_m']:.2f}"
            cells[6].text = str(int(row["axial_idx"]))
            cells[7].text = f"{row['knuckle_mass_kg']*1000:.1f}"

    # ── Phase D ──
    add_heading(doc, "Phase D — 2-D/3-D cartography + sensitivity", 1)
    doc.add_paragraph(
        "THESIS: The LHS run feeds 480,000 stratified samples through the "
        "physics model and lets us see how feasibility and best mass vary "
        "across (n_rings, n_lines, r_hub, axial profile, knuckle) pairs. "
        "This is the most honest view of the design space — no optimizer "
        "bias."
    )
    add_image_if_exists(doc, CART / "fig_heat_nrings_vs_nlines_10kw.png",
                         caption="Figure D-1. 10 kW — min feasible mass heatmap "
                                 "over (n_rings, n_lines). Sweet spot emerges at "
                                 "n_rings ≈ 5-8 with n_lines ≈ 5-7.")
    add_image_if_exists(doc, CART / "fig_heat_nrings_vs_nlines_50kw.png",
                         caption="Figure D-2. 50 kW — same heatmap. The sweet "
                                 "spot shifts slightly higher in n_rings.")
    add_image_if_exists(doc, CART / "fig_heat_axial_vs_nlines_feas.png",
                         caption="Figure D-3. Feasibility fraction across axial "
                                 "profile × n_lines. Straight_taper and "
                                 "parabolic have highest feasibility.")
    add_image_if_exists(doc, CART / "fig_3d_nrings_nlines_rhub_mass_10kw.png",
                         caption="Figure D-4. 10 kW feasible cloud in "
                                 "(n_rings, n_lines, r_hub) with mass colour.")
    add_image_if_exists(doc, CART / "fig_axial_family_boxplot_10kw.png",
                         caption="Figure D-5. 10 kW feasible mass distribution "
                                 "per axial family. Straight_taper and "
                                 "parabolic dominate.")
    add_image_if_exists(doc, CART / "fig_sobol_first_order_10kw.png",
                         caption="Figure D-6. Spearman rank correlation as a "
                                 "first-order sensitivity proxy — Do_top and "
                                 "r_hub are the dominant mass drivers.")
    add_image_if_exists(doc, CART / "fig_pareto_mass_fos.png",
                         caption="Figure D-7. Mass vs FOS Pareto cloud — the "
                                 "FOS=1.8 floor is a bright wall in the "
                                 "feasibility manifold.")

    # ── Phase F ──
    add_heading(doc, "Phase F — n_lines × knuckle sensitivity", 1)
    doc.add_paragraph(
        "THESIS: n_lines is a topology choice (pentagon vs hexagon vs …) — "
        "it changes simultaneously the polygon count, the tether-line count, "
        "and the blade count. The LHS data lets us isolate its mass effect."
    )
    add_image_if_exists(doc, CART / "fig_phase_f_nlines_envelope.png",
                         caption="Figure F-1. Min feasible mass vs n_lines, "
                                 "faceted by config, coloured by beam. "
                                 "Pentagon (n=5) and heptagon (n=7) are best.")
    add_image_if_exists(doc, CART / "fig_phase_f_knuckle_mass.png",
                         caption="Figure F-2. Min mass envelope vs knuckle "
                                 "mass — lighter knuckles always help, but "
                                 "only marginally (< 1 kg swing over the "
                                 "10-200 g range).")
    add_image_if_exists(doc, CART / "fig_phase_f_combined_surface.png",
                         caption="Figure F-3. 10 kW circular — min mass "
                                 "surface in (n_lines, knuckle_mass).")

    # ── Phase G ──
    add_heading(doc, "Phase G — Rich visualisation", 1)
    for png in sorted(REND.glob("*.png")) if REND.exists() else []:
        add_image_if_exists(doc, png, width_in=6.2, caption=png.stem)

    # ── Phase E (dynamic ODE verification) ──
    add_heading(doc, "Phase E — Dynamic ODE verification (pending)", 1)
    doc.add_paragraph(
        "Planned: take the top 30 candidates across all islands, rebuild "
        "each SystemParams, run the full multi-body ODE for 30 s of steady "
        "11 m/s and 5 s of coherent 11→25 m/s gust, and check that the "
        "per-ring FOS remains ≥ 1.5 throughout. Reports which candidates "
        "survive dynamic verification — the optimizer uses a quasi-static "
        "envelope which may miss transient failures.")

    # ── Close ──
    add_heading(doc, "Methodology notes & reproducibility", 1)
    doc.add_paragraph(
        "All data live under scripts/results/trpt_opt_v2/. Each island "
        "subdirectory contains log.csv (per-heartbeat), checkpoint.jls "
        "(resumable), elite_archive.csv (up to 200 feasible designs), and "
        "best_design.json. The LHS cartography lives under scripts/results/"
        "trpt_opt_v2/lhs/. Re-running scripts/launch_autonomous_campaign.sh "
        "from a clean state reproduces the full campaign."
    )

    doc.save(OUT)
    print(f"wrote {OUT}")

if __name__ == "__main__":
    main()

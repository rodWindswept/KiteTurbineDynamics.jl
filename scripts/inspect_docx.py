# scripts/inspect_docx.py
import sys
from docx import Document

def inspect_report(filepath, query_terms):
    print(f"\n==========================================")
    print(f"Inspecting file: {filepath}")
    print(f"==========================================")
    
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"Error opening file: {e}")
        return
        
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")
    
    # Search paragraphs
    print("\n--- Matching Paragraphs ---")
    match_count = 0
    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        if any(term in text for term in query_terms):
            print(f"Paragraph {idx}: {text}")
            match_count += 1
    if match_count == 0:
        print("No paragraphs matched the search terms.")
        
    # Search tables
    print("\n--- Table Headers and Matching Rows ---")
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx}: {len(table.rows)} rows, {len(table.columns)} columns")
        # Print header
        if len(table.rows) > 0:
            header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
            print(f"  Header: {header}")
            
        row_matches = 0
        for r_idx in range(1, len(table.rows)):
            row_cells = table.rows[r_idx].cells
            row_text = " | ".join([cell.text.strip().replace('\n', ' ') for cell in row_cells])
            if any(term in row_text for term in query_terms):
                print(f"  Row {r_idx}: {row_text}")
                row_matches += 1
        if row_matches == 0 and len(table.rows) > 1:
            # print first row data just for structure
            row_cells = table.rows[1].cells
            row_text = " | ".join([cell.text.strip().replace('\n', ' ') for cell in row_cells])
            print(f"  Example Row 1: {row_text}")

if __name__ == "__main__":
    inspect_report("TRPT_Twist_Analysis.docx", ["0.43", "Cp", "peak", "218", "249", "1250"])
    inspect_report("TRPT_Ring_Scalability_Report.docx", ["2333", "tension", "rated", "9.6"])
    inspect_report("TRPT_Lift_Device_Analysis.docx", ["10.27", "hub_z_std", "excursion", "1441"])
